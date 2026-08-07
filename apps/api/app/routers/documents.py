from __future__ import annotations

import asyncio
import json
import os

from fastapi import APIRouter, Depends, File, HTTPException, UploadFile, status
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.db import get_db
from app.deps import get_current_user
from app.models import CoverLetter, Document, Profile, User
from app.schemas import DocumentLibrary, GeneratedDoc
from app.services import llm_service, storage_service

router = APIRouter(prefix="/documents", tags=["documents"])

_ALLOWED = {".pdf", ".docx"}


def _ocr_pdf(path: str) -> str:
    """
    OCR fallback for scanned PDFs: rasterize each page at 300 DPI with pdf2image
    (Poppler) and run Tesseract via pytesseract. No-op (returns "") if either
    optional dependency or the system binaries are unavailable.
    """
    try:
        import pdf2image  # type: ignore
        import pytesseract  # type: ignore

        images = pdf2image.convert_from_path(path, dpi=300)
        return "\n".join(pytesseract.image_to_string(img) for img in images)
    except Exception:
        return ""


def _extract_text(path: str) -> str:
    """
    Read text from an uploaded CV. Tries pdfplumber / python-docx if installed;
    if a PDF yields almost no text (a scanned document) it falls back to OCR;
    otherwise reads the raw bytes as UTF-8 (works for the demo's text uploads).
    Kept dependency-light so the API runs without the heavy parsing stack.
    """
    ext = os.path.splitext(path)[1].lower()
    try:
        if ext == ".pdf":
            import pdfplumber  # type: ignore

            with pdfplumber.open(path) as pdf:
                text = "\n".join(page.extract_text() or "" for page in pdf.pages)
            # < 100 chars indicates a scanned/image PDF — try OCR.
            if len(text.strip()) < 100:
                ocr = _ocr_pdf(path)
                if ocr.strip():
                    return ocr
            return text
        if ext == ".docx":
            import docx  # type: ignore

            return "\n".join(p.text for p in docx.Document(path).paragraphs)
    except Exception:
        pass
    try:
        with open(path, "rb") as fh:
            return fh.read().decode("utf-8", errors="ignore")
    except Exception:
        return ""


@router.post("/upload")
async def upload(
    file: UploadFile = File(...),
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
) -> dict[str, str]:
    ext = os.path.splitext(file.filename or "")[1].lower()
    if ext not in _ALLOWED:
        raise HTTPException(status.HTTP_400_BAD_REQUEST, detail="Only PDF or DOCX allowed")

    # Store via the storage layer — local disk, or S3/R2 when configured. The
    # returned reference (path or object key) is what we persist on the row.
    ref = await storage_service.save(await file.read(), file.filename or f"cv{ext}")

    doc = Document(user_id=user.id, filename=file.filename or ref, path=ref, kind="cv")
    db.add(doc)
    await db.commit()
    await db.refresh(doc)
    return {"documentId": doc.id}


@router.get("/parse")
async def parse(
    documentId: str,
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
) -> StreamingResponse:
    """
    Server-Sent Events stream of parsing progress. Emits one event per stage,
    then a final event carrying the structured profile. Matches the frontend's
    fetch-streaming consumer (which sends the Authorization header).
    """
    doc = await db.get(Document, documentId)
    if not doc or doc.user_id != user.id:
        raise HTTPException(status.HTTP_404_NOT_FOUND, detail="Document not found")

    async def event_stream():
        stages = [
            ("read", "Reading document…"),
            ("work", "Extracting work history…"),
            ("skills", "Identifying skills…"),
            ("education", "Structuring education…"),
            ("contact", "Collecting contact details…"),
        ]
        for stage, message in stages:
            await asyncio.sleep(0.5)
            yield f"data: {json.dumps({'stage': stage, 'message': message, 'done': False})}\n\n"

        # Resolve to a local file (downloads from S3/R2 if that's the backend).
        async with storage_service.local_path(doc.path) as local:
            text = _extract_text(local)
        profile = await llm_service.extract_profile(text)
        doc.parsed = profile
        # Also persist into the user's Profile so the imported CV surfaces across
        # the app immediately — even if the review/confirm step is skipped or the
        # final PUT /profiles/me never runs. The review step still overrides this
        # with any user edits.
        result = await db.execute(select(Profile).where(Profile.user_id == user.id))
        prof = result.scalar_one_or_none()
        if prof is None:
            prof = Profile(user_id=user.id, data=profile, version=int(profile.get("version", 1)))
            db.add(prof)
        else:
            prof.data = profile
            prof.version = int(profile.get("version", 1))
        await db.commit()
        payload = {"stage": "complete", "message": "Done", "done": True, "profile": profile}
        yield f"data: {json.dumps(payload)}\n\n"

    return StreamingResponse(
        event_stream(),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


class ParseTextInput(BaseModel):
    text: str


async def _persist_profile(db: AsyncSession, user_id: str, profile: dict) -> None:
    result = await db.execute(select(Profile).where(Profile.user_id == user_id))
    prof = result.scalar_one_or_none()
    if prof is None:
        prof = Profile(user_id=user_id, data=profile, version=int(profile.get("version", 1)))
        db.add(prof)
    else:
        prof.data = profile
        prof.version = int(profile.get("version", 1))
    await db.commit()


# ---- Source CVs: keep several, choose which one drives the profile ----------
# Client 24.07: "que le permita subir 2 o 3 CV diferentes, y en base a eso esa persona
# pueda seleccionar con que CV vaya probando las funcionalidades de filtros ATS y demás".
# Every upload already creates its own Document(kind="cv"); these endpoints expose that
# history and let the user switch which one the profile (and therefore every AI tool) is
# built from. The choice lives in preferences.activeCvId (prod has no migrations).


class SourceCvOut(BaseModel):
    id: str
    filename: str
    createdAt: str
    active: bool


@router.get("/cvs", response_model=list[SourceCvOut])
async def list_source_cvs(
    user: User = Depends(get_current_user), db: AsyncSession = Depends(get_db)
) -> list[SourceCvOut]:
    """Every CV the user has uploaded, newest first, flagging the active one."""
    rows = (
        await db.execute(
            select(Document)
            .where(Document.user_id == user.id, Document.kind == "cv")
            .order_by(Document.created_at.desc())
        )
    ).scalars().all()
    active_id = (user.preferences or {}).get("activeCvId")
    # With nothing chosen yet the newest upload is the de-facto source of the profile.
    if not active_id and rows:
        active_id = rows[0].id
    return [
        SourceCvOut(
            id=d.id,
            filename=d.filename,
            createdAt=d.created_at.isoformat(),
            active=(d.id == active_id),
        )
        for d in rows
    ]


@router.post("/cvs/{doc_id}/activate")
async def activate_source_cv(
    doc_id: str,
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
) -> dict[str, object]:
    """Make this CV the active one: re-parse it into the profile so every tool (ATS,
    optimizer, job matching) works from it."""
    doc = await db.get(Document, doc_id)
    if not doc or doc.user_id != user.id or doc.kind != "cv":
        raise HTTPException(status.HTTP_404_NOT_FOUND, detail="CV not found")

    text = _extract_text(doc.path)
    if not text.strip():
        raise HTTPException(
            status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="Could not read any text from that file.",
        )
    profile = await llm_service.extract_profile(text)
    await _persist_profile(db, user.id, profile)

    prefs = dict(user.preferences or {})
    prefs["activeCvId"] = doc.id
    user.preferences = prefs
    await db.commit()
    return {"ok": True, "profile": profile}


@router.delete("/cvs/{doc_id}")
async def delete_source_cv(
    doc_id: str,
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
) -> dict[str, bool]:
    """Remove an uploaded CV from the list. The parsed profile is left untouched."""
    doc = await db.get(Document, doc_id)
    if not doc or doc.user_id != user.id or doc.kind != "cv":
        raise HTTPException(status.HTTP_404_NOT_FOUND, detail="CV not found")
    await db.delete(doc)
    prefs = dict(user.preferences or {})
    if prefs.get("activeCvId") == doc_id:
        prefs.pop("activeCvId", None)
        user.preferences = prefs
    await db.commit()
    return {"ok": True}


@router.get("/library", response_model=DocumentLibrary)
async def library(
    user: User = Depends(get_current_user), db: AsyncSession = Depends(get_db)
) -> DocumentLibrary:
    """Generated documents: optimized CVs + cover letters."""
    cv_rows = (
        await db.execute(
            select(Document)
            .where(Document.user_id == user.id, Document.kind == "optimized_cv")
            .order_by(Document.created_at.desc())
        )
    ).scalars().all()
    letter_rows = (
        await db.execute(
            select(CoverLetter)
            .where(CoverLetter.user_id == user.id)
            .order_by(CoverLetter.created_at.desc())
        )
    ).scalars().all()
    cvs = [
        GeneratedDoc(
            id=d.id,
            title=d.filename,
            createdAt=d.created_at,
            preview=((d.parsed or {}).get("text", ""))[:240],
            atsScore=(d.parsed or {}).get("ats"),
        )
        for d in cv_rows
    ]
    letters = [
        GeneratedDoc(
            id=letter.id,
            title=f"Cover letter · {letter.tone}",
            createdAt=letter.created_at,
            preview=(letter.text or "")[:240],
        )
        for letter in letter_rows
    ]
    return DocumentLibrary(cvs=cvs, letters=letters)


@router.get("/generated/{doc_id}")
async def get_generated(
    doc_id: str, user: User = Depends(get_current_user), db: AsyncSession = Depends(get_db)
) -> dict:
    """Full text of a generated document (optimized CV or cover letter)."""
    d = await db.get(Document, doc_id)
    if d and d.user_id == user.id and d.kind == "optimized_cv":
        p = d.parsed or {}
        return {"id": d.id, "title": d.filename, "text": p.get("text", ""), "atsScore": p.get("ats")}
    letter = await db.get(CoverLetter, doc_id)
    if letter and letter.user_id == user.id:
        return {"id": letter.id, "title": f"Cover letter · {letter.tone}", "text": letter.text, "atsScore": None}
    raise HTTPException(status.HTTP_404_NOT_FOUND, detail="Not found")


@router.post("/parse-text")
async def parse_text(
    body: ParseTextInput,
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
) -> dict:
    """Parse a pasted CV (plain text) into the profile — the alternative AI source
    to file upload. Persists straight into Profile.data."""
    profile = await llm_service.extract_profile(body.text or "")
    await _persist_profile(db, user.id, profile)
    return {"profile": profile}
